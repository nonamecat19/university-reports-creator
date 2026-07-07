"""Template ingestion (FR-TPL-08..11): docx bytes → TemplateModel dict +
diagnostics.
"""

from __future__ import annotations

import io
from typing import Any

import docx

from .pm import paragraphs_to_pm_doc
from .placeholders import WELL_KNOWN_FIELDS, field_label, field_type, find_placeholders, find_unclosed
from .safety import UnsafeDocxError, check_zip_safe, has_macros

EMU_PER_TWIP = 635


class TemplateParseError(ValueError):
    """Raised for FR-TPL-10 "errors" — the upload is rejected."""

    def __init__(self, message: str, location: str = ""):
        super().__init__(message)
        self.location = location


def _diag(severity: str, message: str, location: str = "") -> dict:
    return {"severity": severity, "message": message, "location": location}


def _scan_field_placeholders(
    text: str, location: str, fields: dict[str, dict], diagnostics: list[dict]
) -> tuple[bool, bool]:
    """Registers scalar-field placeholder occurrences found in `text`.
    Returns (saw_bibliography, saw_toc)."""
    saw_bibliography = False
    saw_toc = False

    if find_unclosed(text):
        diagnostics.append(_diag("warning", f"placeholder-like text did not parse near {location!r}", location))

    for _match, placeholder in find_placeholders(text):
        if placeholder.kind == "bibliography":
            saw_bibliography = True
        elif placeholder.kind == "toc":
            saw_toc = True
        elif placeholder.kind == "field":
            name = placeholder.name or ""
            existing = fields.get(name)
            if existing is None:
                fields[name] = {
                    "name": name,
                    "label": placeholder.label or field_label(name),
                    "type": field_type(name),
                    "required": placeholder.required,
                    "default": placeholder.default or "",
                    "occurrences": [location],
                }
            else:
                existing["required"] = existing["required"] or placeholder.required
                if location not in existing["occurrences"]:
                    existing["occurrences"].append(location)
    return saw_bibliography, saw_toc


def _scan_body(paragraphs: list[Any], fields: dict[str, dict], diagnostics: list[dict]) -> tuple[list[dict], bool, bool]:
    sections: list[dict] = []
    section_ids_seen: set[str] = set()
    open_section: dict | None = None
    saw_bibliography = False
    saw_toc = False

    for para_index, paragraph in enumerate(paragraphs):
        text = paragraph.text
        if not text:
            if open_section is not None:
                open_section["_paragraphs"].append(paragraph)
            continue

        matches = find_placeholders(text)
        section_marker = next(
            (p for _m, p in matches if p.kind in ("section_start", "section_end")), None
        )

        if section_marker is None:
            bib, toc = _scan_field_placeholders(text, f"body:{para_index}", fields, diagnostics)
            saw_bibliography = saw_bibliography or bib
            saw_toc = saw_toc or toc
            if open_section is not None:
                open_section["_paragraphs"].append(paragraph)
            continue

        stripped = text.strip()
        sole = len(matches) == 1 and stripped == matches[0][0].group(0)
        if not sole:
            diagnostics.append(
                _diag("warning", "section marker should be alone in its paragraph", f"body:{para_index}")
            )

        if section_marker.kind == "section_start":
            section_id = section_marker.name or f"section_{len(sections)}"
            if section_id in section_ids_seen:
                raise TemplateParseError(f"duplicate section id {section_id!r}", f"body:{para_index}")
            if open_section is not None:
                raise TemplateParseError(
                    f"section {open_section['id']!r} not closed before {section_id!r} starts",
                    f"body:{para_index}",
                )
            section_ids_seen.add(section_id)
            open_section = {
                "id": section_id,
                "label": section_marker.label,
                "required": section_marker.required,
                "kind": section_marker.section_kind or "chapter",
                "min_words": section_marker.min_words,
                "order": len(sections),
                "_paragraphs": [],
            }
        else:  # section_end
            if open_section is None:
                raise TemplateParseError("`{{/section}}` without a matching `{{#section:...}}`", f"body:{para_index}")
            sections.append(open_section)
            open_section = None

    if open_section is not None:
        raise TemplateParseError(f"section {open_section['id']!r} was never closed with `{{{{/section}}}}`")

    return sections, saw_bibliography, saw_toc


_ALWAYS_RELEVANT_STYLES = {
    "Normal",
    "Caption",
    "List Bullet",
    "List Number",
    *(f"Heading {i}" for i in range(1, 7)),
}


def _used_style_names(document: Any) -> set[str]:
    used: set[str] = set()
    for paragraph in document.paragraphs:
        if paragraph.style is not None:
            used.add(paragraph.style.name)
    for section in document.sections:
        for paragraph in (*section.header.paragraphs, *section.footer.paragraphs):
            if paragraph.style is not None:
                used.add(paragraph.style.name)
    return used


def _style_map(document: Any) -> dict[str, dict]:
    """Only styles the export translator or the template body actually cares
    about (FR-TPL-08 "named styles ... relevant to content mapping") — a
    fresh docx ships ~150 unused theme styles that would otherwise drown out
    the useful ones."""
    relevant = _ALWAYS_RELEVANT_STYLES | _used_style_names(document)

    style_map: dict[str, dict] = {}
    for style in document.styles:
        if style.type is None or style.name is None or style.name not in relevant:
            continue
        try:
            font = style.font
        except (AttributeError, KeyError):
            continue
        style_map[style.name] = {
            "type": str(style.type),
            "font_name": font.name if font else None,
            "size_pt": font.size.pt if font and font.size else None,
            "bold": font.bold if font else None,
            "italic": font.italic if font else None,
        }
    return style_map


def _page_setup(document: Any) -> dict:
    section = document.sections[0]
    return {
        "page_width_twips": round(section.page_width / EMU_PER_TWIP) if section.page_width else None,
        "page_height_twips": round(section.page_height / EMU_PER_TWIP) if section.page_height else None,
        "margin_top_twips": round(section.top_margin / EMU_PER_TWIP) if section.top_margin else None,
        "margin_bottom_twips": round(section.bottom_margin / EMU_PER_TWIP) if section.bottom_margin else None,
        "margin_left_twips": round(section.left_margin / EMU_PER_TWIP) if section.left_margin else None,
        "margin_right_twips": round(section.right_margin / EMU_PER_TWIP) if section.right_margin else None,
        "orientation": "landscape" if section.orientation == 1 else "portrait",
    }


def parse_template(docx_bytes: bytes, max_decompressed_bytes: int) -> tuple[dict, list[dict]]:
    diagnostics: list[dict] = []

    try:
        check_zip_safe(docx_bytes, max_decompressed_bytes)
    except UnsafeDocxError as exc:
        raise TemplateParseError(str(exc)) from exc

    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise TemplateParseError(f"could not open docx (corrupt or password-protected): {exc}") from exc

    if has_macros(docx_bytes):
        diagnostics.append(_diag("warning", "template contains a macro (vbaProject.bin); stripped at export"))

    fields: dict[str, dict] = {}
    sections, saw_bib_body, saw_toc_body = _scan_body(document.paragraphs, fields, diagnostics)

    saw_bibliography = saw_bib_body
    saw_toc = saw_toc_body
    for section in document.sections:
        for para_index, paragraph in enumerate(section.header.paragraphs):
            bib, toc = _scan_field_placeholders(paragraph.text, f"header:{para_index}", fields, diagnostics)
            saw_bibliography = saw_bibliography or bib
            saw_toc = saw_toc or toc
        for para_index, paragraph in enumerate(section.footer.paragraphs):
            bib, toc = _scan_field_placeholders(paragraph.text, f"footer:{para_index}", fields, diagnostics)
            saw_bibliography = saw_bibliography or bib
            saw_toc = saw_toc or toc

    if not saw_bibliography:
        diagnostics.append(_diag("warning", "no {{bibliography}} marker found in the template"))
    if not sections:
        diagnostics.append(_diag("warning", "no {{#section}} regions defined; document degenerates to metadata-only"))

    for name in WELL_KNOWN_FIELDS:
        if name in fields:
            fields[name]["type"] = field_type(name)

    model_sections = [
        {
            "id": s["id"],
            "label": s["label"],
            "order": s["order"],
            "required": s["required"],
            "kind": s["kind"],
            "min_words": s["min_words"],
            "example_content": paragraphs_to_pm_doc(s["_paragraphs"]),
        }
        for s in sections
    ]

    model = {
        "fields": list(fields.values()),
        "sections": model_sections,
        "style_map": _style_map(document),
        "page_setup": _page_setup(document),
        "numbering": [],
        "warnings": [d["message"] for d in diagnostics if d["severity"] == "warning"],
        "has_toc_marker": saw_toc,
        "has_bibliography_marker": saw_bibliography,
    }
    return model, diagnostics
