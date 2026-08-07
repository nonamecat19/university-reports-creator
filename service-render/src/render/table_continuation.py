"""«Продовження таблиці N» continuation captions (FR-TBL-09).

ДСТУ 3008:2015 practice is that a table spilling onto the next page is
introduced there by «Продовження таблиці N». OOXML has no field for this —
Word only knows how to *repeat header rows* (`w:tblHeader`, the default
`repeat_header` strategy, FR-TBL-08). So the `continuation_caption` strategy
has to know where the page breaks actually fall, which only a layout engine
can tell us.

The pass therefore works like this:

1. take the rendered docx and produce a **probe** copy in which every table row
   carries a tiny white marker run (`⟦t0r3⟧`) — 1 pt and white, so it changes
   layout by essentially nothing but still comes out of PDF text extraction;
2. convert the probe to PDF with LibreOffice and read which page each marker
   landed on;
3. a row whose page is greater than the previous row's page starts a
   continuation — split the real table there, repeat its header rows, and put a
   «Продовження таблиці N» paragraph in between.

This is **best-effort by construction**, exactly as FR-TBL-09 says: the breaks
are the ones LibreOffice computed for this file, and any later edit in Word can
move them. The export surfaces that as a warning.
"""

from __future__ import annotations

import copy
import io
import re
from dataclasses import dataclass
from typing import Any, Callable

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from . import numbering as num

# `⟦`/`⟧` are outside anything a student would type, so the marker cannot be
# confused with real cell text when it is read back out of the PDF.
_MARKER_RE = re.compile(r"⟦t(\d+)r(\d+)⟧")

PdfRenderer = Callable[[bytes], bytes]


@dataclass
class ContinuationPlan:
    """Row indices at which each table starts a new page, per table index."""

    breaks: dict[int, list[int]]


def _marker(table_index: int, row_index: int) -> str:
    return f"⟦t{table_index}r{row_index}⟧"


def build_probe_docx(docx_bytes: bytes) -> bytes:
    """Returns a copy of the document with a layout-neutral marker in every
    table row, for the pagination pass."""
    document = docx.Document(io.BytesIO(docx_bytes))

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cell = row.cells[0]
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = paragraph.add_run(_marker(table_index, row_index))
            run.font.size = Pt(1)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def page_breaks_from_pdf(pdf_bytes: bytes) -> ContinuationPlan:
    """Maps the probe markers to page numbers and returns, per table, the row
    indices that begin a new page."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(pdf_bytes))

    # (table, row) -> first page the marker appears on.
    row_pages: dict[tuple[int, int], int] = {}
    for page_number, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - a page we cannot read simply yields no markers
            continue
        for match in _MARKER_RE.finditer(text):
            key = (int(match.group(1)), int(match.group(2)))
            row_pages.setdefault(key, page_number)

    breaks: dict[int, list[int]] = {}
    for (table_index, row_index), page in sorted(row_pages.items()):
        previous = row_pages.get((table_index, row_index - 1))
        if previous is not None and page > previous:
            breaks.setdefault(table_index, []).append(row_index)

    return ContinuationPlan(breaks=breaks)


def _header_row_count(table: Any) -> int:
    """Rows marked `w:tblHeader` (FR-TBL-08) — repeated at the top of every
    continuation so a split table stays readable."""
    count = 0
    for row in table.rows:
        tr_pr = row._tr.find(qn("w:trPr"))
        if tr_pr is None or tr_pr.find(qn("w:tblHeader")) is None:
            break
        count += 1
    return count


def _split_table(table: Any, break_rows: list[int], number: str) -> int:
    """Splits one table at the given row indices, inserting a continuation
    caption before each part. Returns the number of continuations created."""
    tbl = table._tbl
    header_count = _header_row_count(table)

    created = 0
    # Split back-to-front: cutting the tail off first leaves the earlier break
    # points at the same indices in what remains of the original table.
    for row_index in sorted(break_rows, reverse=True):
        rows = tbl.findall(qn("w:tr"))
        if row_index <= header_count or row_index >= len(rows):
            # A break inside (or right after) the repeated header would produce
            # a continuation whose only content is that header.
            continue

        continuation = copy.deepcopy(tbl)
        for tr in continuation.findall(qn("w:tr")):
            continuation.remove(tr)
        for i in range(header_count):
            continuation.append(copy.deepcopy(rows[i]))
        for tr in rows[row_index:]:
            tbl.remove(tr)
            continuation.append(tr)

        caption = OxmlElement("w:p")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = num.table_continuation_caption(number)
        run.append(text)
        caption.append(run)
        # keepNext binds the caption to the table it introduces, the same way
        # the main caption is bound (FR-TBL-08).
        p_pr = OxmlElement("w:pPr")
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)
        caption.insert(0, p_pr)

        tbl.addnext(continuation)
        tbl.addnext(caption)
        created += 1

    return created


def apply_continuation_captions(
    document: Any,
    plan: ContinuationPlan,
    numbers_by_table_index: list[str],
) -> int:
    """Splits every paginated table and inserts its continuation captions.
    Returns how many continuation captions were added."""
    tables = list(document.tables)
    added = 0

    # Later tables are unaffected by earlier splits only if we work backwards:
    # splitting inserts new tables into the body.
    for table_index in sorted(plan.breaks, reverse=True):
        if table_index >= len(tables):
            continue
        number = (
            numbers_by_table_index[table_index]
            if table_index < len(numbers_by_table_index)
            else ""
        )
        if not number:
            # An unnumbered table has nothing to say in «Продовження таблиці N».
            continue
        added += _split_table(tables[table_index], plan.breaks[table_index], number)

    return added


def add_continuation_captions(
    docx_bytes: bytes,
    document: Any,
    numbers_by_table_index: list[str],
    render_pdf: PdfRenderer,
) -> tuple[int, list[dict]]:
    """Runs the whole pass against an already-rendered document.

    `docx_bytes` is that document as saved (the probe is built from it);
    `document` is the live python-docx object the splits are applied to, so the
    caller can simply save it again afterwards.
    """
    warnings: list[dict] = []
    if not document.tables:
        return 0, warnings

    try:
        pdf_bytes = render_pdf(build_probe_docx(docx_bytes))
    except Exception as exc:  # noqa: BLE001 - the strategy degrades, the export does not fail
        warnings.append({
            "severity": "warning",
            "message": f"«Продовження таблиці» pagination pass failed, falling back to repeated header rows ({exc})",
            "location": "",
        })
        return 0, warnings

    plan = page_breaks_from_pdf(pdf_bytes)
    added = apply_continuation_captions(document, plan, numbers_by_table_index)

    if added:
        warnings.append({
            "severity": "info",
            "message": (
                f"inserted {added} «Продовження таблиці» caption(s) at the page breaks LibreOffice computed; "
                "editing the file in Word can move those breaks"
            ),
            "location": "",
        })
    return added, warnings
