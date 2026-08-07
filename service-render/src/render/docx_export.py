"""Export translator (FR-EXP-01..03): merges document content into the
original template docx. Deterministic — same input bytes in, same output
bytes out (export determinism is an acceptance criterion, FR-EXP export
suite).
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any

import docx
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.table import Table
from docx.text.paragraph import Paragraph

from . import numbering as num
from .bibliography import DEFAULT_STYLE, render_entries
from .docx_xml import (
    insert_toc_field,
    mark_header_row,
    new_paragraph_after,
    relocate_after,
    remove_paragraphs,
    set_cant_split,
    set_run_marks,
    set_update_fields_on_open,
)
from .omml import UnsupportedLatex, latex_to_omml
from .placeholders import find_placeholders
from .revisions import CommentsBuilder, RevisionIds, should_emit_text, suggestion_mark, wrap_revision
from .table_continuation import PdfRenderer, add_continuation_captions

PX_TO_EMU = 9525


@dataclass
class RenderSectionInput:
    id: str
    template_section_id: str
    title: str
    kind: str  # "chapter" | "appendix"
    order: int
    content: dict[str, Any] | None


@dataclass
class RenderComment:
    """One comment to export as a native Word comment (FR-REV-13)."""

    section_id: str
    block_id: str
    body: str
    author_id: str
    timestamp: str = ""


@dataclass
class TranslateContext:
    """Everything the translator needs beyond the node being translated:
    tracked-changes state, comment anchors, and the collected warnings.

    Threaded through translation rather than held globally so `render_docx`
    stays a pure function of its inputs (export determinism, FR-EXP).
    """

    strategy: str
    ids: RevisionIds
    comments: CommentsBuilder | None
    # block_id -> comments anchored on it, consumed as blocks are emitted.
    comments_by_block: dict[str, list[RenderComment]]
    authors: dict[str, str]
    # source_id -> reference-list number, for in-text `[N]` citations (FR-BIB-05).
    citation_numbers: dict[str, int]
    # target_id -> «рис. 2.1», for cross-references (FR-EDT-04/07).
    reference_labels: dict[str, str]
    # id(w:tbl element) -> table number, for continuation captions (FR-TBL-09).
    table_numbers: dict[int, str]
    # Non-fatal export problems surfaced on the job (FR-EXP-05), appended to
    # as nodes are translated.
    warnings: list[dict]


def _substitute_scalar_fields(paragraph: Paragraph, values: dict[str, str]) -> None:
    runs = paragraph.runs
    if not runs:
        return
    full_text = "".join(r.text for r in runs)
    matches = find_placeholders(full_text)
    if not matches:
        return

    pieces: list[str] = []
    last_end = 0
    changed = False
    for match, placeholder in matches:
        if placeholder.kind != "field":
            continue
        value = values.get(placeholder.name or "", placeholder.default or "")
        pieces.append(full_text[last_end : match.start()])
        pieces.append(value)
        last_end = match.end()
        changed = True

    if not changed:
        return
    pieces.append(full_text[last_end:])
    runs[0].text = "".join(pieces)
    for r in runs[1:]:
        r.text = ""


def _find_marker(paragraphs: list[Paragraph], kind: str) -> Paragraph | None:
    for paragraph in paragraphs:
        for _match, placeholder in find_placeholders(paragraph.text):
            if placeholder.kind == kind:
                return paragraph
    return None


def _find_section_regions(paragraphs: list[Paragraph]) -> dict[str, tuple[int, int]]:
    """Returns {section_id: (start_index, end_index)} over the given
    paragraph list (indices inclusive of the marker paragraphs)."""
    regions: dict[str, tuple[int, int]] = {}
    open_id: str | None = None
    open_start = -1
    for i, paragraph in enumerate(paragraphs):
        for _match, placeholder in find_placeholders(paragraph.text):
            if placeholder.kind == "section_start":
                open_id = placeholder.name
                open_start = i
            elif placeholder.kind == "section_end" and open_id is not None:
                regions[open_id] = (open_start, i)
                open_id = None
    return regions


def _inline_text_runs(paragraph: Paragraph, inline_nodes: list[dict], revisions: TranslateContext) -> None:
    for node in inline_nodes or []:
        node_type = node.get("type")

        if node_type == "citation":
            _emit_citation(paragraph, node, revisions)
            continue

        if node_type == "crossReference":
            _emit_cross_reference(paragraph, node, revisions)
            continue

        if node_type == "formulaInline":
            _emit_formula(paragraph, (node.get("attrs") or {}).get("latex", ""), revisions)
            continue

        if node_type != "text":
            continue

        marks = node.get("marks") or []
        suggestion = suggestion_mark(marks)
        if not should_emit_text(suggestion, revisions.strategy):
            continue

        run = paragraph.add_run(node.get("text", ""))
        set_run_marks(run, marks)
        if suggestion is not None and revisions.strategy == "with_track_changes":
            wrap_revision(run, suggestion, revisions.ids, revisions.authors)


def _emit_citation(paragraph: Paragraph, node: dict[str, Any], revisions: TranslateContext) -> None:
    """Writes an in-text citation as the bracketed reference-list number
    (FR-BIB-05). The number is computed from the reference list, never stored
    in the node, so it always agrees with what the bibliography printed."""
    attrs = node.get("attrs") or {}
    source_id = attrs.get("sourceId", "")
    locator = (attrs.get("locator") or "").strip()
    number = revisions.citation_numbers.get(source_id)

    if number is None:
        # Export validation blocks on orphan citations (FR-BIB-07); if one still
        # reaches here, print the marker the editor showed rather than nothing.
        paragraph.add_run("[?]")
        revisions.warnings.append({
            "severity": "warning",
            "message": "citation refers to a source that is not in the reference list",
            "location": source_id,
        })
        return

    paragraph.add_run(f"[{number}, {locator}]" if locator else f"[{number}]")


def _emit_cross_reference(paragraph: Paragraph, node: dict[str, Any], revisions: TranslateContext) -> None:
    """Writes a cross-reference as its resolved label, e.g. «рис. 2.1»
    (FR-EDT-04/07). Both the number and the wording come from the numbering
    counters, so the exported text matches what the editor displayed.

    Word fields (`REF`) are deliberately not used: the numbers they would track
    are ours, not Word's — the document has no Word bookmarks to point at.
    """
    target_id = (node.get("attrs") or {}).get("targetId", "")
    label = revisions.reference_labels.get(target_id)

    if label is None:
        paragraph.add_run(num.UNRESOLVED_REFERENCE)
        revisions.warnings.append({
            "severity": "warning",
            "message": "cross-reference points at a block that no longer exists",
            "location": target_id,
        })
        return

    paragraph.add_run(label)


def _emit_formula(paragraph: Paragraph, latex: str, revisions: TranslateContext) -> None:
    """Appends the formula as a native Word equation (FR-EDT-06).

    `m:oMath` is a direct child of `w:p`, not of a run, so it is appended to the
    paragraph element itself. Constructs outside the converter's supported
    subset fall back to the literal LaTeX source in Cambria Math plus a warning
    — the student then sees exactly what did not convert instead of a silently
    dropped formula.
    """
    try:
        paragraph._p.append(latex_to_omml(latex))
    except UnsupportedLatex as exc:
        run = paragraph.add_run(latex)
        run.font.name = "Cambria Math"
        revisions.warnings.append({
            "severity": "warning",
            "message": f"formula exported as plain LaTeX source ({exc})",
            "location": latex,
        })


def _translate_node(
    node: dict[str, Any],
    document: Any,
    images: dict[str, bytes],
    numbering_result: num.NumberingResult,
    anchor: Any,
    revisions: TranslateContext,
) -> Any:
    node_type = node.get("type")
    block_id = (node.get("attrs") or {}).get("blockId")
    number = numbering_result.block_numbers.get(block_id or "")

    if node_type == "paragraph":
        para = new_paragraph_after(anchor, document)
        _inline_text_runs(para, node.get("content"), revisions)
        _attach_block_comments(para, block_id, revisions)
        return para._p

    if node_type == "heading":
        level = int((node.get("attrs") or {}).get("level", 2))
        para = new_paragraph_after(anchor, document)
        try:
            para.style = document.styles[f"Heading {level}"]
        except KeyError:
            pass
        if number:
            para.add_run(f"{number} ").bold = True
        _inline_text_runs(para, node.get("content"), revisions)
        _attach_block_comments(para, block_id, revisions)
        return para._p

    if node_type in ("bulletList", "orderedList"):
        style_name = "List Bullet" if node_type == "bulletList" else "List Number"
        current = anchor
        for item in node.get("content") or []:
            for child in item.get("content") or []:
                para = new_paragraph_after(current, document)
                try:
                    para.style = document.styles[style_name]
                except KeyError:
                    pass
                _inline_text_runs(para, child.get("content"), revisions)
                current = para._p
        return current

    if node_type == "table":
        rows = node.get("content") or []
        col_count = max((len(r.get("content") or []) for r in rows), default=1)
        row_count = max(len(rows), 1)

        # Caption goes ABOVE the table (FR-TBL-04), bound with keepNext so it
        # never dangles at a page bottom (FR-TBL-08).
        current = anchor
        if number:
            caption_para = new_paragraph_after(current, document)
            caption_para.add_run(num.table_caption(number, (node.get("attrs") or {}).get("caption", "")))
            caption_para.paragraph_format.keep_with_next = True
            current = caption_para._p

        table = document.add_table(rows=row_count, cols=col_count)
        relocate_after(table._tbl, current)
        # Remembered by element identity (not by position) so the continuation
        # pass can label «Продовження таблиці N» correctly (FR-TBL-09).
        revisions.table_numbers[id(table._tbl)] = number or ""

        header_row_count = 0
        for r_idx, row in enumerate(rows):
            cells = row.get("content") or []
            is_header_row = any(c.get("type") == "tableHeader" for c in cells)
            if is_header_row:
                header_row_count += 1
            for c_idx, cell in enumerate(cells):
                if c_idx >= col_count:
                    break
                docx_cell = table.cell(r_idx, c_idx)
                docx_cell.text = ""
                cell_para = docx_cell.paragraphs[0]
                for block in cell.get("content") or []:
                    if block.get("type") == "paragraph":
                        _inline_text_runs(cell_para, block.get("content"), revisions)

        for i in range(header_row_count):
            mark_header_row(table, i)
        set_cant_split(table)

        return table._tbl

    if node_type == "image":
        object_key = (node.get("attrs") or {}).get("objectKey")
        image_bytes = images.get(object_key) if object_key else None
        current = anchor
        if image_bytes:
            section = document.sections[0]
            text_width_emu = section.page_width - section.left_margin - section.right_margin
            natural_width_px = (node.get("attrs") or {}).get("naturalWidth")
            width = Emu(min(int(natural_width_px) * PX_TO_EMU, text_width_emu)) if natural_width_px else Emu(text_width_emu)
            picture = document.add_picture(io.BytesIO(image_bytes), width=width)
            image_paragraph = document.paragraphs[-1]
            relocate_after(image_paragraph._p, anchor)
            current = image_paragraph._p
        if number:
            caption_para = new_paragraph_after(current, document)
            caption_para.add_run(num.figure_caption(number, (node.get("attrs") or {}).get("caption", "")))
            current = caption_para._p
        return current

    if node_type == "formulaBlock":
        # ДСТУ 3008:2015: the formula sits centred on its own line with the
        # number right-aligned on the same line. Word has no "centre + right"
        # alignment, so the classic recipe is a left-aligned paragraph with a
        # centre tab stop mid-text-width and a right tab stop at the margin.
        para = new_paragraph_after(anchor, document)
        page = document.sections[0]
        text_width = page.page_width - page.left_margin - page.right_margin
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Emu(int(text_width / 2)), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Emu(int(text_width)), WD_TAB_ALIGNMENT.RIGHT)

        para.add_run().add_tab()
        _emit_formula(para, (node.get("attrs") or {}).get("latex", ""), revisions)
        if number:
            number_run = para.add_run()
            number_run.add_tab()
            number_run.add_text(num.formula_number(number))
        _attach_block_comments(para, block_id, revisions)
        return para._p

    if node_type == "codeBlock":
        para = new_paragraph_after(anchor, document)
        text = "".join(n.get("text", "") for n in node.get("content") or [] if n.get("type") == "text")
        run = para.add_run(text)
        run.font.name = "Courier New"
        return para._p

    if node_type == "horizontalRule":
        para = new_paragraph_after(anchor, document)
        run = para.add_run()
        run.add_break(7)  # WD_BREAK.PAGE
        return para._p

    return anchor


def _attach_block_comments(paragraph: Paragraph, block_id: str | None, revisions: TranslateContext) -> None:
    """Anchors every comment on this block, once (FR-REV-13)."""
    if not block_id or revisions.comments is None:
        return
    for comment in revisions.comments_by_block.pop(block_id, []):
        revisions.comments.add(paragraph, comment.body, comment.author_id, comment.timestamp)


def _translate_section_content(
    content: dict[str, Any] | None,
    document: Any,
    images: dict[str, bytes],
    numbering_result: num.NumberingResult,
    anchor: Any,
    revisions: TranslateContext,
) -> Any:
    if not content:
        return anchor
    current = anchor
    for node in content.get("content") or []:
        current = _translate_node(node, document, images, numbering_result, current, revisions)
    return current


def render_docx(
    template_docx: bytes,
    metadata: dict[str, str],
    sections: list[RenderSectionInput],
    sources_csl_json: list[str],
    images: dict[str, bytes],
    numbering_mode: str,
    suggestions_strategy: str = "clean",
    comments: list[RenderComment] | None = None,
    authors: dict[str, str] | None = None,
    table_continuation: str = "repeat_header",
    render_pdf: PdfRenderer | None = None,
    citation_style: str = "",
) -> tuple[bytes, list[dict]]:
    warnings: list[dict] = []
    document = docx.Document(io.BytesIO(template_docx))

    ids = RevisionIds()
    comments_by_block: dict[str, list[RenderComment]] = {}
    for comment in comments or []:
        comments_by_block.setdefault(comment.block_id, []).append(comment)

    # Rendered once, up front: in-text citations need the numbers while sections
    # are translated, and the reference list needs the same entries later.
    bibliography_entries = render_entries(
        sources_csl_json, numbering_mode, citation_style or DEFAULT_STYLE
    )

    # Numbering likewise has to exist before translation starts — cross-references
    # in the first section may point at a figure in the last one.
    numbering_result = num.compute_numbering(
        [num.NumberingInput(id=s.id, kind=s.kind, order=s.order, content=s.content) for s in sections]
    )

    revisions = TranslateContext(
        strategy=suggestions_strategy or "clean",
        ids=ids,
        comments=CommentsBuilder(ids, authors or {}) if comments_by_block else None,
        comments_by_block=comments_by_block,
        authors=authors or {},
        citation_numbers={e["source_id"]: e["number"] for e in bibliography_entries if e["source_id"]},
        reference_labels=num.reference_labels(numbering_result),
        table_numbers={},
        warnings=warnings,
    )

    body_paragraphs = list(document.paragraphs)
    for paragraph in body_paragraphs:
        _substitute_scalar_fields(paragraph, metadata)
    for section in document.sections:
        for paragraph in list(section.header.paragraphs):
            _substitute_scalar_fields(paragraph, metadata)
        for paragraph in list(section.footer.paragraphs):
            _substitute_scalar_fields(paragraph, metadata)

    regions = _find_section_regions(body_paragraphs)
    by_template_id = {s.template_section_id: s for s in sections if s.template_section_id}
    extra_sections = sorted((s for s in sections if not s.template_section_id), key=lambda s: s.order)

    for region_id, (start_idx, end_idx) in regions.items():
        matching = by_template_id.get(region_id)
        region_paragraphs = body_paragraphs[start_idx : end_idx + 1]
        if matching is None:
            warnings.append(
                {"severity": "warning", "message": f"no document section for template region {region_id!r}; left as-is", "location": ""}
            )
            continue
        anchor = region_paragraphs[0]._p
        _translate_section_content(matching.content, document, images, numbering_result, anchor, revisions)
        remove_paragraphs(region_paragraphs)

    if extra_sections:
        tail_anchor = document.add_paragraph()._p
        for section in extra_sections:
            heading_para = new_paragraph_after(tail_anchor, document)
            try:
                heading_para.style = document.styles["Heading 1"]
            except KeyError:
                pass
            label = numbering_result.section_labels.get(section.id, "")
            heading_para.add_run(f"{num.section_label(section.kind, label)}. {section.title}").bold = True
            tail_anchor = heading_para._p
            tail_anchor = _translate_section_content(section.content, document, images, numbering_result, tail_anchor, revisions)

    bib_marker = _find_marker(list(document.paragraphs), "bibliography")
    if bib_marker is not None:
        anchor = bib_marker._p
        for entry in bibliography_entries:
            para = new_paragraph_after(anchor, document)
            para.add_run(f"{entry['number']}. {entry['formatted']}")
            anchor = para._p
        remove_paragraphs([bib_marker])

    toc_marker = _find_marker(list(document.paragraphs), "toc")
    if toc_marker is not None:
        for run in list(toc_marker.runs):
            run.text = ""
        insert_toc_field(toc_marker)

    if revisions.comments is not None:
        revisions.comments.attach(document)
    # A comment whose anchor block is gone (deleted since it was written) has
    # nowhere to land in the file; it stays visible in-app as an orphan.
    for block_id, orphaned in revisions.comments_by_block.items():
        warnings.append({
            "severity": "warning",
            "message": f"{len(orphaned)} comment(s) anchored on a missing block were not exported",
            "location": block_id,
        })

    set_update_fields_on_open(document)

    out = io.BytesIO()
    document.save(out)

    # FR-TBL-09: the `continuation_caption` strategy needs the finished file to
    # know where its own page breaks fall, so it runs as a second pass over the
    # saved bytes and then re-saves.
    if table_continuation == "continuation_caption" and render_pdf is not None:
        added, continuation_warnings = add_continuation_captions(
            out.getvalue(),
            document,
            [revisions.table_numbers.get(id(t._tbl), "") for t in document.tables],
            render_pdf,
        )
        warnings.extend(continuation_warnings)
        if added:
            out = io.BytesIO()
            document.save(out)

    return out.getvalue(), warnings
