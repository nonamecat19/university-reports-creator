"""«Продовження таблиці N» continuation captions (FR-TBL-09).

The pagination itself comes from LibreOffice, which is not available in unit
tests — so the tests drive the pass with a stub PDF renderer and assert on the
two halves that are ours: the probe markers, and the table split.
"""

import io

import docx
import pytest

from render import numbering as num
from render.docx_export import RenderSectionInput, render_docx
from render.table_continuation import (
    ContinuationPlan,
    apply_continuation_captions,
    build_probe_docx,
    page_breaks_from_pdf,
)



def table_docx(rows: int, header: bool = True) -> tuple[bytes, docx.Document]:
    document = docx.Document()
    table = document.add_table(rows=rows, cols=2)
    for r in range(rows):
        for c in range(2):
            table.cell(r, c).text = f"r{r}c{c}"
    if header:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tr_pr.append(docx.oxml.OxmlElement("w:tblHeader"))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), document


class TestProbe:
    def test_every_row_gets_a_marker(self):
        source, _ = table_docx(4)
        probe = docx.Document(io.BytesIO(build_probe_docx(source)))
        table = probe.tables[0]

        for row_index in range(4):
            assert f"⟦t0r{row_index}⟧" in table.cell(row_index, 0).text

    def test_markers_are_invisible_by_size_and_colour(self):
        """The probe must not change layout, or the breaks it measures would
        not be the breaks of the real file."""
        source, _ = table_docx(2)
        probe = docx.Document(io.BytesIO(build_probe_docx(source)))
        run = probe.tables[0].cell(0, 0).paragraphs[0].runs[-1]

        assert run.font.size.pt == 1
        assert str(run.font.color.rgb) == "FFFFFF"

    def test_original_is_not_modified(self):
        source, _ = table_docx(2)
        build_probe_docx(source)
        untouched = docx.Document(io.BytesIO(source))
        assert "⟦" not in untouched.tables[0].cell(0, 0).text


class TestPageBreakDetection:
    def test_break_is_where_the_page_number_increases(self, monkeypatch):
        pages = [
            "⟦t0r0⟧ ⟦t0r1⟧ ⟦t0r2⟧",
            "⟦t0r3⟧ ⟦t0r4⟧",
            "⟦t0r5⟧",
        ]

        class FakePage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class FakeReader:
            def __init__(self, _stream) -> None:
                self.pages = [FakePage(t) for t in pages]

        import pypdf

        monkeypatch.setattr(pypdf, "PdfReader", FakeReader)

        plan = page_breaks_from_pdf(b"")
        assert plan.breaks == {0: [3, 5]}

    def test_a_table_on_one_page_has_no_breaks(self, monkeypatch):
        class FakePage:
            def extract_text(self) -> str:
                return "⟦t0r0⟧ ⟦t0r1⟧"

        class FakeReader:
            def __init__(self, _stream) -> None:
                self.pages = [FakePage()]

        import pypdf

        monkeypatch.setattr(pypdf, "PdfReader", FakeReader)
        assert page_breaks_from_pdf(b"").breaks == {}


class TestSplit:
    def test_split_inserts_a_caption_and_repeats_the_header(self):
        _, document = table_docx(6)

        added = apply_continuation_captions(document, ContinuationPlan(breaks={0: [4]}), ["2.1"])
        assert added == 1
        assert len(document.tables) == 2

        first, second = document.tables
        # Rows 0..3 stay; the header is repeated on top of rows 4..5.
        assert len(first.rows) == 4
        assert len(second.rows) == 3
        assert second.cell(0, 0).text == "r0c0"
        assert second.cell(1, 0).text == "r4c0"

        body_text = [p.text for p in document.paragraphs]
        assert "Продовження таблиці 2.1" in body_text

    def test_multiple_breaks_produce_multiple_continuations(self):
        _, document = table_docx(9)

        added = apply_continuation_captions(document, ContinuationPlan(breaks={0: [3, 6]}), ["1.1"])
        assert added == 2
        assert len(document.tables) == 3
        assert [p.text for p in document.paragraphs].count("Продовження таблиці 1.1") == 2

    def test_break_inside_the_header_is_ignored(self):
        """Splitting at the header row would produce a continuation whose only
        content is the repeated header."""
        _, document = table_docx(4)
        assert apply_continuation_captions(document, ContinuationPlan(breaks={0: [1]}), ["1.1"]) == 0
        assert len(document.tables) == 1

    def test_unnumbered_table_is_left_alone(self):
        _, document = table_docx(6)
        assert apply_continuation_captions(document, ContinuationPlan(breaks={0: [3]}), [""]) == 0
        assert len(document.tables) == 1

    def test_caption_text(self):
        assert num.table_continuation_caption("2.1") == "Продовження таблиці 2.1"


@pytest.fixture
def template_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("{{#section:intro|label=Вступ}}")
    document.add_paragraph("{{/section}}")
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def section_with_table(rows: int) -> RenderSectionInput:
    return RenderSectionInput(
        id="sec-1",
        template_section_id="intro",
        title="Вступ",
        kind="chapter",
        order=0,
        content={
            "type": "doc",
            "content": [{
                "type": "table",
                "attrs": {"blockId": "t1", "caption": "Дані"},
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {"type": "tableHeader" if r == 0 else "tableCell",
                             "content": [{"type": "paragraph", "content": [{"type": "text", "text": f"r{r}c{c}"}]}]}
                            for c in range(2)
                        ],
                    }
                    for r in range(rows)
                ],
            }],
        },
    )


class TestRenderIntegration:
    def test_repeat_header_never_paginates(self, template_bytes):
        """The default strategy must not invoke LibreOffice at all."""
        def fail(_bytes: bytes) -> bytes:
            raise AssertionError("pagination must not run for repeat_header")

        _, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with_table(4)],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
            table_continuation="repeat_header",
            render_pdf=fail,
        )
        assert warnings == []

    def test_pagination_failure_degrades_to_a_warning(self, template_bytes):
        def broken(_bytes: bytes) -> bytes:
            raise RuntimeError("soffice missing")

        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with_table(4)],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
            table_continuation="continuation_caption",
            render_pdf=broken,
        )
        assert any("falling back to repeated header rows" in w["message"] for w in warnings)
        # The export still produced a usable file.
        assert docx.Document(io.BytesIO(out)).tables

    def test_continuation_caption_uses_the_table_number(self, template_bytes, monkeypatch):
        class FakePage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class FakeReader:
            def __init__(self, _stream) -> None:
                self.pages = [FakePage("⟦t0r0⟧ ⟦t0r1⟧ ⟦t0r2⟧"), FakePage("⟦t0r3⟧")]

        import pypdf

        monkeypatch.setattr(pypdf, "PdfReader", FakeReader)

        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with_table(4)],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
            table_continuation="continuation_caption",
            render_pdf=lambda _b: b"%PDF-fake",
        )

        document = docx.Document(io.BytesIO(out))
        assert any("Продовження таблиці 1.1" in p.text for p in document.paragraphs)
        assert any("computed" in w["message"] for w in warnings)
