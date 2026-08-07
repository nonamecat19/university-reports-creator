"""Footnote export (FR-EDT-04): the editor's inline footnote node has to become
a real `w:footnote` in its own part, not a superscript digit in the body."""

import io
import zipfile

import docx
import pytest
from lxml import etree

from render.docx_export import RenderSectionInput, render_docx
from render.footnotes import (
    CONTINUATION_SEPARATOR_ID,
    FOOTNOTES_CONTENT_TYPE,
    FOOTNOTES_RELTYPE,
    SEPARATOR_ID,
)

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@pytest.fixture
def template_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("{{#section:intro|label=Вступ}}")
    document.add_paragraph("приклад")
    document.add_paragraph("{{/section}}")
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def section_with(*footnote_texts: str) -> RenderSectionInput:
    content = [{"type": "text", "text": "Твердження"}]
    for text in footnote_texts:
        content.append({"type": "footnote", "attrs": {"text": text}})
    return RenderSectionInput(
        id="sec-1",
        template_section_id="intro",
        title="Вступ",
        kind="chapter",
        order=0,
        content={
            "type": "doc",
            "content": [
                {"type": "paragraph", "attrs": {"blockId": "block-1"}, "content": content}
            ],
        },
    )


def render(template_bytes: bytes, *footnote_texts: str):
    return render_docx(
        template_docx=template_bytes,
        metadata={},
        sections=[section_with(*footnote_texts)],
        sources_csl_json=[],
        images={},
        numbering_mode="by_order",
    )


def part(docx_bytes: bytes, name: str) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return zf.read(name).decode("utf-8")


def names(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return zf.namelist()


class TestFootnotePart:
    def test_document_without_footnotes_gets_no_part(self):
        """An empty footnotes part is legal but makes Word offer an empty
        footnote area, so it must not be written at all."""
        document = docx.Document()
        document.add_paragraph("{{#section:intro|label=Вступ}}")
        document.add_paragraph("{{/section}}")
        out = io.BytesIO()
        document.save(out)

        result, _ = render(out.getvalue())
        assert "word/footnotes.xml" not in names(result)

    def test_footnote_text_lands_in_the_part(self, template_bytes):
        result, _ = render(template_bytes, "Джерело: власні спостереження")

        assert "Джерело: власні спостереження" in part(result, "word/footnotes.xml")

    def test_reference_is_written_into_the_body(self, template_bytes):
        result, _ = render(template_bytes, "примітка")

        assert "footnoteReference" in part(result, "word/document.xml")

    def test_note_text_stays_out_of_the_body(self, template_bytes):
        """A note rendered inline instead of in the part would show up twice in
        Word — once in the text, once at the bottom of the page."""
        result, _ = render(template_bytes, "примітка")

        assert "примітка" not in part(result, "word/document.xml")

    def test_separators_come_first(self, template_bytes):
        result, _ = render(template_bytes, "примітка")
        root = etree.fromstring(part(result, "word/footnotes.xml").encode("utf-8"))

        ids = [f.get(f"{W_NS}id") for f in root.findall(f"{W_NS}footnote")]
        assert ids[:2] == [str(SEPARATOR_ID), str(CONTINUATION_SEPARATOR_ID)]

    def test_each_note_gets_its_own_id(self, template_bytes):
        result, _ = render(template_bytes, "перша", "друга", "третя")
        root = etree.fromstring(part(result, "word/footnotes.xml").encode("utf-8"))

        content_ids = [
            f.get(f"{W_NS}id")
            for f in root.findall(f"{W_NS}footnote")
            if f.get(f"{W_NS}type") is None
        ]
        assert content_ids == ["1", "2", "3"]

    def test_body_references_match_the_note_ids(self, template_bytes):
        result, _ = render(template_bytes, "перша", "друга")
        body = etree.fromstring(part(result, "word/document.xml").encode("utf-8"))

        referenced = [
            ref.get(f"{W_NS}id") for ref in body.iter(f"{W_NS}footnoteReference")
        ]
        assert referenced == ["1", "2"]

    def test_note_carries_its_own_reference_mark(self, template_bytes):
        """Without `w:footnoteRef` the note prints with no number beside it."""
        result, _ = render(template_bytes, "примітка")

        assert "footnoteRef" in part(result, "word/footnotes.xml")

    def test_part_is_declared_in_content_types(self, template_bytes):
        result, _ = render(template_bytes, "примітка")

        assert FOOTNOTES_CONTENT_TYPE in part(result, "[Content_Types].xml")

    def test_part_is_related_to_the_document(self, template_bytes):
        result, _ = render(template_bytes, "примітка")

        assert FOOTNOTES_RELTYPE in part(result, "word/_rels/document.xml.rels")

    def test_leading_space_is_preserved(self, template_bytes):
        """The space between the mark and the text would otherwise collapse."""
        result, _ = render(template_bytes, "примітка")
        root = etree.fromstring(part(result, "word/footnotes.xml").encode("utf-8"))

        texts = [t for t in root.iter(f"{W_NS}t") if t.text and "примітка" in t.text]
        assert texts[0].get("{http://www.w3.org/XML/1998/namespace}space") == "preserve"


class TestEmptyFootnote:
    def test_empty_note_is_not_exported(self, template_bytes):
        result, _ = render(template_bytes, "   ")

        assert "word/footnotes.xml" not in names(result)

    def test_empty_note_warns(self, template_bytes):
        _, warnings = render(template_bytes, "   ")

        assert any("empty footnote" in w["message"] for w in warnings)

    def test_remaining_notes_still_export(self, template_bytes):
        result, _ = render(template_bytes, "   ", "справжня")
        root = etree.fromstring(part(result, "word/footnotes.xml").encode("utf-8"))

        content_ids = [
            f.get(f"{W_NS}id")
            for f in root.findall(f"{W_NS}footnote")
            if f.get(f"{W_NS}type") is None
        ]
        assert content_ids == ["1"]


class TestTemplateWithExistingFootnotes:
    def test_template_part_is_replaced_not_duplicated(self):
        """A template that already carries footnotes must not end up with two
        conflicting footnotes relationships."""
        document = docx.Document()
        document.add_paragraph("{{#section:intro|label=Вступ}}")
        document.add_paragraph("{{/section}}")

        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        stale = Part(
            PackURI("/word/footnotes.xml"),
            FOOTNOTES_CONTENT_TYPE,
            b'<?xml version="1.0"?><w:footnotes xmlns:w="'
            + W_NS[1:-1].encode()
            + b'"></w:footnotes>',
            document.part.package,
        )
        document.part.relate_to(stale, FOOTNOTES_RELTYPE)
        out = io.BytesIO()
        document.save(out)

        result, _ = render(out.getvalue(), "примітка")
        rels = part(result, "word/_rels/document.xml.rels")

        assert rels.count(FOOTNOTES_RELTYPE) == 1
        assert "примітка" in part(result, "word/footnotes.xml")
