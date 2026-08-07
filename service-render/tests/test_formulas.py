"""Export of the derived inline/block nodes: formulas as OMML (FR-EDT-06),
citations as `[N]` (FR-BIB-05) and cross-references as resolved labels
(FR-EDT-04/07). What these nodes have in common is that they store only a
reference and render from the live counters, so the tests check that the docx
text matches what the editor computed."""

import io

import docx
import pytest
from lxml import etree

from render import numbering as num
from render.docx_export import RenderSectionInput, render_docx
from render.omml import UnsupportedLatex, latex_to_omml

M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def omml_xml(latex: str) -> str:
    return etree.tostring(latex_to_omml(latex), encoding="unicode")


class TestLatexToOmml:
    def test_plain_text_becomes_math_runs(self):
        xml = omml_xml("E = mc")
        assert xml.count("<m:r>") == 4
        assert ">E<" in xml

    def test_superscript(self):
        assert "<m:sSup>" in omml_xml("x^2")

    def test_subscript(self):
        assert "<m:sSub>" in omml_xml("x_i")

    @pytest.mark.parametrize("source", ["x_i^2", "x^2_i"])
    def test_sub_and_sup_merge_into_one_element(self, source):
        """Both orders must produce a single sSubSup, not nested scripts —
        otherwise Word renders the second script attached to the first."""
        xml = omml_xml(source)
        assert "<m:sSubSup>" in xml
        assert "<m:sSup>" not in xml.replace("<m:sSubSup>", "")

    def test_fraction(self):
        xml = omml_xml(r"\frac{a+b}{c}")
        assert "<m:f>" in xml
        assert "<m:num>" in xml and "<m:den>" in xml

    def test_square_root_hides_the_degree(self):
        xml = omml_xml(r"\sqrt{x}")
        assert '<m:degHide m:val="1"' in xml

    def test_nth_root_keeps_the_degree(self):
        xml = omml_xml(r"\sqrt[3]{x}")
        assert 'degHide' in xml and 'val="0"' in xml
        assert ">3<" in xml

    def test_sum_with_limits(self):
        xml = omml_xml(r"\sum_{i=1}^{n} a_i")
        assert "<m:nary>" in xml
        assert "∑" in xml
        assert 'val="undOvr"' in xml

    def test_integral_uses_side_limits(self):
        xml = omml_xml(r"\int_0^1 x dx")
        assert "∫" in xml
        assert 'val="subSup"' in xml

    def test_limit_renders_as_a_function_with_a_lower_limit(self):
        xml = omml_xml(r"\lim_{x \to 0} \frac{1}{x}")
        assert "<m:limLow>" in xml
        assert ">lim<" in xml

    def test_left_right_delimiters(self):
        xml = omml_xml(r"\left(\frac{x}{y}\right)")
        assert "<m:d>" in xml
        assert 'begChr' in xml and 'endChr' in xml

    def test_greek_and_operators_map_to_unicode(self):
        xml = omml_xml(r"\alpha \times \beta \leq \infty")
        for glyph in ("α", "×", "β", "≤", "∞"):
            assert glyph in xml

    def test_text_command_is_upright(self):
        xml = omml_xml(r"\text{Обсяг} = \pi r^2")
        assert "<m:nor/>" in xml
        assert "Обсяг" in xml

    def test_function_names_are_upright(self):
        xml = omml_xml(r"\sin x")
        assert "<m:nor/>" in xml
        assert ">sin<" in xml

    @pytest.mark.parametrize(
        "source",
        [
            "",
            "   ",
            r"\begin{matrix}a & b\end{matrix}",
            r"a \over b",
            r"x \\ y",
            r"\unknowncommand{x}",
            r"\frac{a}{b",
            r"\left( x",
            "^2",
        ],
    )
    def test_unsupported_sources_raise(self, source):
        with pytest.raises(UnsupportedLatex):
            latex_to_omml(source)

    def test_conversion_is_deterministic(self):
        """Export determinism (FR-EXP): identical source, identical bytes."""
        assert omml_xml(r"\frac{\sum_{i=1}^{n} x_i}{n}") == omml_xml(r"\frac{\sum_{i=1}^{n} x_i}{n}")


class TestFormulaNumbering:
    def test_formulas_number_per_chapter(self):
        result = num.compute_numbering([
            num.NumberingInput(
                id="s1",
                kind="chapter",
                order=0,
                content={
                    "type": "doc",
                    "content": [
                        {"type": "formulaBlock", "attrs": {"blockId": "f1", "latex": "a"}},
                        {"type": "paragraph", "attrs": {"blockId": "p1"}},
                        {"type": "formulaBlock", "attrs": {"blockId": "f2", "latex": "b"}},
                    ],
                },
            ),
            num.NumberingInput(
                id="s2",
                kind="chapter",
                order=1,
                content={
                    "type": "doc",
                    "content": [{"type": "formulaBlock", "attrs": {"blockId": "f3", "latex": "c"}}],
                },
            ),
        ])
        assert result.block_numbers["f1"] == "1.1"
        assert result.block_numbers["f2"] == "1.2"
        assert result.block_numbers["f3"] == "2.1"

    def test_appendix_formulas_use_the_appendix_letter(self):
        result = num.compute_numbering([
            num.NumberingInput(
                id="a1",
                kind="appendix",
                order=0,
                content={
                    "type": "doc",
                    "content": [{"type": "formulaBlock", "attrs": {"blockId": "f1", "latex": "a"}}],
                },
            )
        ])
        assert result.block_numbers["f1"] == "А.1"

    def test_number_format(self):
        assert num.formula_number("2.3") == "(2.3)"


@pytest.fixture
def template_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("{{#section:intro|label=Вступ}}")
    document.add_paragraph("приклад")
    document.add_paragraph("{{/section}}")
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def section_with(content_nodes: list[dict]) -> RenderSectionInput:
    return RenderSectionInput(
        id="sec-1",
        template_section_id="intro",
        title="Вступ",
        kind="chapter",
        order=0,
        content={"type": "doc", "content": content_nodes},
    )


class TestFormulaExport:
    def test_block_formula_emits_native_omml_with_a_right_aligned_number(self, template_bytes):
        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with([
                {"type": "formulaBlock", "attrs": {"blockId": "f1", "latex": r"E = mc^2"}}
            ])],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
        )
        assert warnings == []

        document = docx.Document(io.BytesIO(out))
        body = document.element.body
        assert body.findall(f".//{M_NS}oMath")

        formula_paragraph = next(
            p for p in document.paragraphs if p._p.findall(f".//{M_NS}oMath")
        )
        assert "(1.1)" in formula_paragraph.text
        alignments = [t.alignment for t in formula_paragraph.paragraph_format.tab_stops]
        assert len(alignments) == 2  # centre for the formula, right for the number

    def test_inline_formula_lands_inside_the_paragraph(self, template_bytes):
        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with([
                {
                    "type": "paragraph",
                    "attrs": {"blockId": "p1"},
                    "content": [
                        {"type": "text", "text": "де "},
                        {"type": "formulaInline", "attrs": {"latex": "x_i"}},
                        {"type": "text", "text": " — значення."},
                    ],
                }
            ])],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
        )
        assert warnings == []
        document = docx.Document(io.BytesIO(out))
        paragraph = next(p for p in document.paragraphs if "значення" in p.text)
        assert paragraph._p.findall(f".//{M_NS}oMath")

    def test_unsupported_formula_falls_back_to_source_with_a_warning(self, template_bytes):
        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with([
                {
                    "type": "formulaBlock",
                    "attrs": {"blockId": "f1", "latex": r"\begin{matrix}a & b\end{matrix}"},
                }
            ])],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
        )
        assert len(warnings) == 1
        assert "plain LaTeX source" in warnings[0]["message"]

        document = docx.Document(io.BytesIO(out))
        assert any("matrix" in p.text for p in document.paragraphs)


class TestCitationExport:
    def test_in_text_citation_renders_the_reference_number(self, template_bytes):
        source = '{"id": "src-1", "type": "book", "title": "Аналіз даних", "issued": {"date-parts": [[2020]]}}'
        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with([
                {
                    "type": "paragraph",
                    "attrs": {"blockId": "p1"},
                    "content": [
                        {"type": "text", "text": "як показано в "},
                        {"type": "citation", "attrs": {"sourceId": "src-1", "locator": "с. 45"}},
                    ],
                }
            ])],
            sources_csl_json=[source],
            images={},
            numbering_mode="by_order",
        )
        assert warnings == []
        document = docx.Document(io.BytesIO(out))
        assert any("[1, с. 45]" in p.text for p in document.paragraphs)

    def test_orphan_citation_warns_and_prints_a_marker(self, template_bytes):
        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with([
                {
                    "type": "paragraph",
                    "attrs": {"blockId": "p1"},
                    "content": [{"type": "citation", "attrs": {"sourceId": "gone"}}],
                }
            ])],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
        )
        assert any("not in the reference list" in w["message"] for w in warnings)
        document = docx.Document(io.BytesIO(out))
        assert any("[?]" in p.text for p in document.paragraphs)


class TestCrossReferenceExport:
    """FR-EDT-04/07: references carry only a target id; wording and number are
    resolved from the same counters the editor showed."""

    def test_reference_labels_cover_every_numbered_target(self):
        result = num.compute_numbering([
            num.NumberingInput(
                id="s1",
                kind="chapter",
                order=0,
                content={
                    "type": "doc",
                    "content": [
                        {"type": "heading", "attrs": {"blockId": "h1", "level": 2}},
                        {"type": "image", "attrs": {"blockId": "i1"}},
                        {"type": "table", "attrs": {"blockId": "t1"}},
                        {"type": "formulaBlock", "attrs": {"blockId": "f1", "latex": "x"}},
                    ],
                },
            ),
            num.NumberingInput(id="a1", kind="appendix", order=1, content=None),
        ])
        labels = num.reference_labels(result)

        assert labels["h1"] == "розд. 1.1"
        assert labels["i1"] == "рис. 1.1"
        assert labels["t1"] == "табл. 1.1"
        assert labels["f1"] == "(1.1)"
        assert labels["s1"] == "розділ 1"
        assert labels["a1"] == "додаток А"

    def test_reference_exports_the_resolved_label(self, template_bytes):
        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with([
                {"type": "image", "attrs": {"blockId": "i1", "caption": "Схема"}},
                {
                    "type": "paragraph",
                    "attrs": {"blockId": "p1"},
                    "content": [
                        {"type": "text", "text": "див. "},
                        {"type": "crossReference", "attrs": {"targetId": "i1"}},
                    ],
                },
            ])],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
        )
        assert warnings == []
        document = docx.Document(io.BytesIO(out))
        assert any("див. рис. 1.1" in p.text for p in document.paragraphs)

    def test_reference_to_a_later_section_resolves(self, template_bytes):
        """A reference in section 1 may point at a figure in section 2, so
        numbering has to be complete before translation starts."""
        template = docx.Document()
        template.add_paragraph("{{#section:intro|label=Вступ}}")
        template.add_paragraph("{{/section}}")
        template.add_paragraph("{{#section:body|label=Основна частина}}")
        template.add_paragraph("{{/section}}")
        buffer = io.BytesIO()
        template.save(buffer)

        sections = [
            RenderSectionInput(
                id="s1",
                template_section_id="intro",
                title="Вступ",
                kind="chapter",
                order=0,
                content={
                    "type": "doc",
                    "content": [{
                        "type": "paragraph",
                        "attrs": {"blockId": "p1"},
                        "content": [{"type": "crossReference", "attrs": {"targetId": "i1"}}],
                    }],
                },
            ),
            RenderSectionInput(
                id="s2",
                template_section_id="body",
                title="Основна частина",
                kind="chapter",
                order=1,
                content={
                    "type": "doc",
                    "content": [{"type": "image", "attrs": {"blockId": "i1"}}],
                },
            ),
        ]

        out, warnings = render_docx(
            template_docx=buffer.getvalue(),
            metadata={},
            sections=sections,
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
        )
        assert warnings == []
        document = docx.Document(io.BytesIO(out))
        assert any("рис. 2.1" in p.text for p in document.paragraphs)

    def test_orphan_reference_warns_and_prints_a_marker(self, template_bytes):
        out, warnings = render_docx(
            template_docx=template_bytes,
            metadata={},
            sections=[section_with([
                {
                    "type": "paragraph",
                    "attrs": {"blockId": "p1"},
                    "content": [{"type": "crossReference", "attrs": {"targetId": "gone"}}],
                }
            ])],
            sources_csl_json=[],
            images={},
            numbering_mode="by_order",
        )
        assert any("no longer exists" in w["message"] for w in warnings)
        document = docx.Document(io.BytesIO(out))
        assert any(num.UNRESOLVED_REFERENCE in p.text for p in document.paragraphs)
