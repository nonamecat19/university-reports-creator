"""Track-changes and comment export (FR-REV-13, FR-EXP-04)."""

import io
import zipfile

import docx
import pytest

from render.docx_export import RenderComment, RenderSectionInput, render_docx
from render.revisions import should_emit_text, suggestion_mark

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@pytest.fixture
def template_bytes() -> bytes:
    """A minimal template with one editable section region (FR-TPL-05)."""
    document = docx.Document()
    document.add_paragraph("{{#section:intro|label=Вступ}}")
    document.add_paragraph("приклад")
    document.add_paragraph("{{/section}}")
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def section_with_suggestions() -> RenderSectionInput:
    return RenderSectionInput(
        id="sec-1",
        template_section_id="intro",
        title="Вступ",
        kind="chapter",
        order=0,
        content={
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "attrs": {"blockId": "block-1"},
                    "content": [
                        {"type": "text", "text": "звичайний "},
                        {
                            "type": "text",
                            "text": "доданий",
                            "marks": [
                                {
                                    "type": "suggestionInsert",
                                    "attrs": {
                                        "suggestionId": "sug-1",
                                        "authorId": "user-1",
                                        "timestamp": "2026-03-14T10:00:00Z",
                                    },
                                }
                            ],
                        },
                        {
                            "type": "text",
                            "text": "видалений",
                            "marks": [
                                {
                                    "type": "suggestionDelete",
                                    "attrs": {"suggestionId": "sug-2", "authorId": "user-1"},
                                }
                            ],
                        },
                    ],
                }
            ],
        },
    )


def render(template_bytes: bytes, strategy: str, comments=None, include=True):
    return render_docx(
        template_docx=template_bytes,
        metadata={},
        sections=[section_with_suggestions()],
        sources_csl_json=[],
        images={},
        numbering_mode="by_order",
        suggestions_strategy=strategy,
        comments=comments if include else [],
        authors={"user-1": "Петренко Петро"},
    )


def body_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def all_text(docx_bytes: bytes) -> str:
    return "\n".join(p.text for p in docx.Document(io.BytesIO(docx_bytes)).paragraphs)


class TestStrategySelection:
    def test_no_mark_always_emits(self):
        for strategy in ("clean", "all_accepted", "with_track_changes"):
            assert should_emit_text(None, strategy) is True

    def test_accepting_drops_deletions(self):
        mark = {"type": "suggestionDelete"}
        assert should_emit_text(mark, "all_accepted") is False
        assert should_emit_text({"type": "suggestionInsert"}, "all_accepted") is True

    def test_clean_drops_insertions(self):
        assert should_emit_text({"type": "suggestionInsert"}, "clean") is False
        assert should_emit_text({"type": "suggestionDelete"}, "clean") is True

    def test_suggestion_mark_ignores_formatting_marks(self):
        assert suggestion_mark([{"type": "bold"}, {"type": "italic"}]) is None
        assert suggestion_mark([{"type": "bold"}, {"type": "suggestionInsert"}])["type"] == "suggestionInsert"


class TestTrackChangesExport:
    def test_with_track_changes_emits_native_revisions(self, template_bytes):
        docx_bytes, _ = render(template_bytes, "with_track_changes")
        xml = body_xml(docx_bytes)

        assert f"{W_NS}ins" in xml.replace("w:ins", f"{W_NS}ins") or "<w:ins " in xml
        assert "<w:del " in xml
        # Deleted text must be w:delText, else Word drops it silently.
        assert "<w:delText" in xml
        assert 'w:author="Петренко Петро"' in xml
        assert 'w:date="2026-03-14T10:00:00Z"' in xml

    def test_clean_export_has_no_revisions_and_rejects_insertions(self, template_bytes):
        docx_bytes, _ = render(template_bytes, "clean")
        xml = body_xml(docx_bytes)
        text = all_text(docx_bytes)

        assert "<w:ins " not in xml
        assert "<w:del " not in xml
        assert "доданий" not in text
        assert "видалений" in text

    def test_all_accepted_applies_insertions_and_drops_deletions(self, template_bytes):
        docx_bytes, _ = render(template_bytes, "all_accepted")
        text = all_text(docx_bytes)

        assert "доданий" in text
        assert "видалений" not in text
        assert "<w:ins " not in body_xml(docx_bytes)


class TestCommentExport:
    def test_comments_produce_a_comments_part_with_anchors(self, template_bytes):
        comment = RenderComment(
            section_id="sec-1", block_id="block-1", body="Уточніть формулювання",
            author_id="user-1", timestamp="2026-03-14T12:00:00Z",
        )
        docx_bytes, warnings = render(template_bytes, "clean", comments=[comment])

        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            names = zf.namelist()
            assert "word/comments.xml" in names
            comments_xml = zf.read("word/comments.xml").decode("utf-8")

        assert "Уточніть формулювання" in comments_xml
        assert 'w:author="Петренко Петро"' in comments_xml
        assert "<w:commentRangeStart" in body_xml(docx_bytes)
        assert "<w:commentReference" in body_xml(docx_bytes)
        assert warnings == []

    def test_no_comments_means_no_comments_part(self, template_bytes):
        docx_bytes, _ = render(template_bytes, "clean", comments=[])
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            assert "word/comments.xml" not in zf.namelist()

    def test_comment_on_missing_block_warns_instead_of_failing(self, template_bytes):
        comment = RenderComment(section_id="sec-1", block_id="gone", body="Осиротілий", author_id="user-1")
        _, warnings = render(template_bytes, "clean", comments=[comment])

        assert len(warnings) == 1
        assert warnings[0]["location"] == "gone"
